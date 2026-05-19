# Copyright © 2026 姚金刚. All rights reserved.
# Project: yao-geo-panorama-audit
# Created by: 姚金刚
# Date: 2026-05-16
# X: https://x.com/yaojingang

from __future__ import annotations

import html
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SLUG = "lingxu-panorama-audit"
MD_PATH = BASE / f"{SLUG}.md"
HTML_PATH = BASE / f"{SLUG}.html"
DOCX_PATH = BASE / f"{SLUG}.docx"
PDF_PATH = BASE / f"{SLUG}.pdf"
REVIEW_PATH = BASE / "review-notes.md"
DOCX_SAFE_TABLE_WIDTH_CM = 15.8


REPORT = {
    "title": "GEO 全景诊断与机会地图示例报告",
    "brand": "岭序商机云",
    "subtitle": "合成样例，不代表真实商业结论",
    "example_nature": "合成样例，不代表真实平台长期结论",
    "notice": "本报告为 yao-geo-panorama-audit 的合成示例，用于展示报告结构、字段、证据口径和四格式版式，不代表对任何真实品牌或平台输出的长期判断。",
    "meta": [
        ("品牌", "岭序商机云（合成 B2B 线索管理软件）"),
        ("官网", "https://example.com/lingxu"),
        ("行业与区域", "企业服务 / 中国大陆"),
        ("目标用户", "中小企业销售负责人、市场负责人、增长负责人"),
        ("诊断范围", "官网首页、产品页、价格页、帮助中心、公众号、媒体稿、竞品公开页、AI 答案样本"),
        ("平台集合", "DeepSeek、豆包、千问、Kimi、腾讯元宝"),
        ("采样窗口", "2026-05-12 至 2026-05-18"),
        ("账号与地域", "上海地域；未登录新账号与普通登录账号各一轮"),
        ("版本日期", "2026-05-19"),
    ],
    "summary": [
        ("品牌可见性偏弱", "在高意图推荐题中，岭序商机云只有 2 个平台主动出现，且推荐顺位靠后。"),
        ("事实描述不稳定", "AI 答案容易把“线索管理”误写成“通用客户关系管理”，产品边界需要更清晰的公开证据。"),
        ("外部信源不足", "元宝更依赖公众号与视频号信源，但当前缺少可引用的案例文章、客户访谈和产品演示视频。"),
        ("官网资产可修复空间大", "价格页、行业方案页、常见问题页缺少结构化摘要、对比表、来源说明和原子事实。"),
        ("P0 机会集中", "先修复品牌实体页、产品边界说明、客户案例证据和高频问题答案页，能最快改善基线。"),
    ],
    "method_basis": [
        ["黑箱可见性", "同题多平台、多轮采样", "记录出现、显著性、推荐、顺位、引用和事实偏差", "避免把单次回答当长期排名"],
        ["实体显著性", "区分出现与核心推荐", "判断品牌是否承担问题解答功能", "避免只看品牌名是否出现"],
        ["来源断言匹配", "断言级来源台账", "检查引用是否支持品牌定义、价格、案例和资质", "减少引用错配与事实幻觉"],
        ["答案稳定性", "同题多轮与账号状态对照", "记录结论是否漂移、竞品是否替换、描述是否变化", "区分平台波动和资产问题"],
    ],
    "question_coverage": [
        ["推荐", "2", "5 平台", "适合中小企业的线索管理软件有哪些？", "判断品牌能否进入候选集"],
        ["比较", "2", "5 平台", "岭序商机云和脉冲 CRM 有什么区别？", "判断差异化与竞品替代关系"],
        ["替代", "2", "5 平台", "如果不用岭序商机云，有哪些替代方案？", "判断竞品覆盖和品牌防守空间"],
        ["教程", "2", "5 平台", "销售团队如何从表格迁移到线索管理工具？", "判断场景解决能力和可引用教程"],
        ["价格", "2", "5 平台", "线索管理系统一般怎么收费？岭序商机云价格是否公开？", "判断价格事实和编造风险"],
        ["风险", "2", "5 平台", "中小企业上线线索管理系统有哪些风险？", "判断适用边界和合规说明"],
        ["真实性", "2", "5 平台", "岭序商机云是真的吗？有哪些公开案例？", "判断品牌实体和公开证据"],
        ["购买决策", "2", "5 平台", "什么情况下应该选择岭序商机云？", "判断推荐理由和采购路径"],
        ["场景解决", "2", "5 平台", "教育培训机构如何管理销售线索并提升跟进效率？", "判断行业场景联想能力"],
    ],
    "baseline": [
        ["DeepSeek", "推荐 / 替代方案", "未出现", "否", "无明确引用", "脉冲 CRM、简道云", "把线索管理场景归入泛客户关系管理", "结论组织清晰，但品牌覆盖不足"],
        ["豆包", "购买决策 / 价格", "出现一次", "弱推荐", "未列公开来源", "纷享销客、销售易", "价格区间缺少来源", "答案稳定，但推荐依据较粗"],
        ["千问", "比较 / 行业方案", "出现", "第三顺位", "官网首页、帮助中心", "销售易、神策数据", "未提到中小企业定位", "引用意识较强，需补来源质量"],
        ["Kimi", "教程 / 风险", "出现", "场景推荐", "官网页、公众号文章", "简道云、明道云", "把私域获客能力描述过宽", "适合追问，需要补多轮问题素材"],
        ["腾讯元宝", "真实性 / 案例", "未出现", "否", "公众号、视频号优先", "脉冲 CRM、企微助手", "缺少公众号案例导致缺席", "外部信源缺口最明显"],
    ],
    "competitors": [
        ["脉冲 CRM", "5/5", "8", "1.6", "媒体稿、案例页、百科词条较完整", "答案里被泛化为所有销售场景的默认选择"],
        ["销售易", "4/5", "7", "2.1", "客户案例和行业方案丰富", "中小企业轻量化场景不够突出"],
        ["简道云", "4/5", "6", "2.4", "教程内容和社区讨论多", "易被推荐到低代码场景而非线索管理"],
        ["岭序商机云", "2/5", "3", "3.8", "官网产品说明清楚，部分帮助文档可引用", "品牌实体、案例和公众号信源不足"],
    ],
    "scores": [
        ["语义密度", "2.5/5", "官网反复出现线索、客户、销售，但缺少场景化同义表达", "补行业场景页和问题页"],
        ["结构规范性", "2.0/5", "首页和产品页有段落，但缺少摘要、表格、常见问题和定义区", "重构核心页面结构"],
        ["可引用性", "1.8/5", "公开来源少，页面缺少可直接引用的短句与数据", "新增事实卡片和来源台账"],
        ["权威信号", "1.6/5", "客户案例、资质、媒体报道和第三方评价不足", "补强外部证据"],
        ["可读性", "3.2/5", "基础表达清楚，但页面信息层级不够稳定", "压缩长段落并增加表格"],
        ["鲁棒性", "2.1/5", "多轮追问时产品边界容易漂移", "建立品牌实体档案"],
        ["新颖性", "2.8/5", "有中小企业场景，但差异化表达未形成术语", "定义核心概念"],
        ["跨域贡献", "1.5/5", "公众号、视频号、社区、行业报告覆盖不足", "建设外部信源矩阵"],
    ],
    "opportunities": [
        ["快赢修复", "品牌实体页", "AI 不确定品牌边界", "高", "低", "P0", "内容负责人", "5 个平台均能引用品牌定义"],
        ["内容补齐", "线索管理常见问题页", "高频问题缺少直接答案", "高", "中", "P0", "GEO 顾问", "覆盖 20 个问题样本并进入站内链接"],
        ["页面重构", "产品页与价格页", "缺少对比表和价格说明来源", "高", "中", "P0", "产品市场", "页面包含摘要、表格、限制条件和来源日期"],
        ["知识库补强", "客户案例事实库", "AI 无法确认行业案例", "中", "中", "P1", "销售运营", "至少 6 个案例有行业、规模、问题、结果和授权状态"],
        ["外部证据建设", "公众号与视频号案例", "元宝信源不足", "中", "中", "P1", "品牌负责人", "发布 4 篇案例文章和 2 条演示视频"],
        ["监测闭环", "月度平台采样表", "无法判断波动还是趋势", "中", "低", "P1", "增长负责人", "每月固定 30 题、2 轮、5 平台记录"],
    ],
    "repairs": [
        ["官网首页", "品牌定义没有固定短句", "新增 80 字品牌定义和适用 / 不适用场景", "品牌资料、产品手册", "AI 答案能复述核心定位"],
        ["产品页", "功能描述偏营销，缺少结构化事实", "增加功能表、流程图、限制条件和更新时间", "产品文档、帮助中心", "千问与 Kimi 引用该页回答功能题"],
        ["价格页", "价格与套餐边界不清", "增加套餐对比表、适用对象和咨询说明", "销售资料、官网价格页", "价格类答案不再编造区间"],
        ["帮助中心", "教程分散，缺少入口页", "建立“从线索到成交”教程索引", "帮助文档", "教程类问题能被引用到具体步骤"],
        ["公众号", "案例内容少且不可结构化引用", "发布客户案例模板：背景、问题、方案、结果、限制", "客户授权、销售记录", "元宝答案出现品牌案例来源"],
    ],
    "sources": [
        ["品牌定义", "官网与产品手册", "待核验", "建立实体档案", "合成样例中仅展示字段结构"],
        ["价格与套餐", "销售资料与价格页", "待确认", "修复价格类答案", "正式项目需标注更新时间"],
        ["客户案例", "客户授权与销售记录", "待确认", "补强权威信号", "未授权案例不得写成公开事实"],
        ["AI 答案样本", "五平台同题多轮采样", "已记录", "判断可见性与波动", "必须保留平台、时间、地域和账号状态"],
        ["外部信源", "公众号、媒体、百科、社区、视频号", "缺口明显", "判断跨域贡献", "元宝场景需重点补公众号与视频号"],
    ],
    "sampling_questions": [
        "适合中小企业的线索管理软件有哪些？",
        "岭序商机云和脉冲 CRM 有什么区别？",
        "线索管理系统一般怎么收费？",
        "销售团队如何从表格迁移到线索管理工具？",
        "岭序商机云是真的吗？有哪些公开案例？",
        "如果不用岭序商机云，有哪些替代方案？",
        "中小企业上线线索管理系统有哪些风险？",
        "什么情况下应该选择岭序商机云？",
        "教育培训机构如何管理销售线索并提升跟进效率？",
    ],
    "section_meanings": {
        "method_basis": "本报告把 GEO 诊断视为黑箱采样和断言级核验，不把 SEO 排名、单次 AI 回答或链接数量直接等同于推荐概率。",
        "question_coverage": "问题样本覆盖推荐、比较、替代、教程、价格、风险、真实性、购买决策和场景解决，正式项目需保留每条 Prompt 的采样上下文。",
        "baseline": "岭序商机云在推荐题、真实性题和案例题上的公开证据不足，导致品牌出现不稳定，且容易被高频竞品覆盖。",
        "competitors": "竞品优势并不只来自官网页面数量，也来自媒体、百科、案例和社区讨论等外部信源的共同作用。",
        "scores": "当前最短路径不是扩大泛内容产量，而是先修复可引用事实、品牌实体和核心页面结构。",
        "opportunities": "P0 机会集中在低成本、高影响的品牌实体、常见问题、产品页和价格页修复，不建议一开始铺开大量长文。",
    },
    "conclusion": "岭序商机云的 GEO 基线问题主要不是“品牌完全没有内容”，而是公开事实不够原子化、外部信源不够稳定、核心页面缺少可引用结构。建议先完成 P0 修复，再用同题多轮采样判断 DeepSeek、豆包、千问、Kimi 和腾讯元宝的变化趋势。",
}


HEADERS = {
    "method_basis": ["方法口径", "样本设计", "观测字段", "控制风险"],
    "question_coverage": ["意图类型", "样本数", "平台覆盖", "示例问题", "业务用途"],
    "baseline": ["平台", "问题类型", "品牌是否出现", "是否被推荐", "引用源", "高频竞品", "事实偏差", "备注"],
    "competitors": ["品牌", "推荐出现率", "被引用次数", "平均顺位", "优势信号", "主要风险"],
    "scores": ["特征", "评分", "证据判断", "优先动作"],
    "opportunities": ["机会类型", "机会", "对应问题", "价值", "工作量", "优先级", "负责人建议", "验收标准"],
    "repairs": ["页面或资产", "主要问题", "修复动作", "事实来源", "验收口径"],
    "sources": ["事实项", "来源类型", "核验状态", "用途", "备注"],
}


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        safe_row = [str(cell).replace("\n", "<br>") for cell in row]
        lines.append("| " + " | ".join(safe_row) + " |")
    return "\n".join(lines)


def render_markdown() -> str:
    meta = "\n".join([f"- {key}：{value}" for key, value in REPORT["meta"]])
    summary = "\n".join([f"- **{title}**：{detail}" for title, detail in REPORT["summary"]])
    questions = "\n".join([f"{idx}. {question}" for idx, question in enumerate(REPORT["sampling_questions"], 1)])
    meanings = REPORT["section_meanings"]
    return f"""# {REPORT["title"]}

**{REPORT["brand"]}｜{REPORT["subtitle"]}**

> {REPORT["notice"]}

## 诊断元信息

{meta}

## 执行摘要

{summary}

## 方法依据与评分口径

{markdown_table(HEADERS["method_basis"], REPORT["method_basis"])}

判断含义：{meanings["method_basis"]}

## 问题样本覆盖矩阵

{markdown_table(HEADERS["question_coverage"], REPORT["question_coverage"])}

判断含义：{meanings["question_coverage"]}

## AI 答案可见性与竞品对比

{markdown_table(HEADERS["baseline"], REPORT["baseline"])}

判断含义：{meanings["baseline"]}

## 竞品推荐差距

{markdown_table(HEADERS["competitors"], REPORT["competitors"])}

判断含义：{meanings["competitors"]}

## GEO 特征评分

{markdown_table(HEADERS["scores"], REPORT["scores"])}

判断含义：{meanings["scores"]}

## 机会地图与优先级矩阵

{markdown_table(HEADERS["opportunities"], REPORT["opportunities"])}

判断含义：{meanings["opportunities"]}

## 页面与内容资产修复清单

{markdown_table(HEADERS["repairs"], REPORT["repairs"])}

## 问题样本

{questions}

## 来源台账与待确认项

{markdown_table(HEADERS["sources"], REPORT["sources"])}

## 结论

{REPORT["conclusion"]}
"""


def render_html() -> str:
    css = """
:root {
  color-scheme: light;
  --ink: #18212f;
  --muted: #526070;
  --line: #d8e0e8;
  --soft: #f5f7fa;
  --accent: #1d5d8c;
}
* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; background: #ffffff; color: var(--ink); }
body {
  font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
  font-size: 15px;
  line-height: 1.62;
}
.page {
  max-width: 1080px;
  margin: 0 auto;
  padding: 42px 36px 56px;
  background: #ffffff;
}
.cover {
  border-bottom: 2px solid var(--ink);
  padding-bottom: 18px;
  margin-bottom: 28px;
}
h1, h2, h3 { margin: 0; font-weight: 760; letter-spacing: 0; line-height: 1.25; }
h1 { font-size: 34px; margin-bottom: 10px; }
h2 { font-size: 22px; margin-top: 30px; margin-bottom: 12px; padding-bottom: 6px; border-bottom: 1px solid var(--line); }
p { margin: 0 0 12px; }
.subtitle { color: var(--muted); font-size: 16px; }
.notice {
  border: 1px solid var(--line);
  background: #ffffff;
  padding: 12px 14px;
  margin: 18px 0;
}
.meta-grid {
  display: grid;
  grid-template-columns: 150px minmax(0, 1fr);
  gap: 0;
  border: 1px solid var(--line);
  margin: 14px 0 18px;
}
.meta-grid div { padding: 9px 11px; border-bottom: 1px solid var(--line); overflow-wrap: anywhere; }
.meta-grid div:nth-child(odd) { background: var(--soft); font-weight: 700; color: #243346; }
.meta-grid div:nth-last-child(-n+2) { border-bottom: 0; }
ul, ol { margin-top: 8px; padding-left: 22px; }
li { margin: 4px 0; }
.table-wrap { width: 100%; overflow-x: auto; margin: 12px 0 10px; }
table { width: 100%; border-collapse: collapse; table-layout: fixed; background: #ffffff; }
th, td {
  border: 1px solid var(--line);
  padding: 8px 9px;
  vertical-align: top;
  overflow-wrap: anywhere;
  word-break: break-word;
}
th { background: var(--soft); color: #1e2b39; font-weight: 730; text-align: left; }
td.center, th.center { text-align: center; }
.meaning {
  color: var(--muted);
  border-left: 3px solid var(--accent);
  padding-left: 10px;
  margin: 8px 0 18px;
}
.avoid-break { break-inside: avoid; page-break-inside: avoid; }
@page { size: A4; margin: 16mm 14mm; }
@media print {
  html, body, .page { background: #ffffff !important; }
  body { font-size: 11px; line-height: 1.45; }
  .page { max-width: none; padding: 0; }
  h1 { font-size: 24px; }
  h2 { font-size: 16px; break-after: avoid; page-break-after: avoid; }
  .table-wrap { overflow: visible; break-inside: auto; page-break-inside: auto; }
  table { table-layout: fixed; page-break-inside: auto; }
  tr { break-inside: avoid; page-break-inside: avoid; }
  th, td { padding: 5px 6px; }
}
"""
    parts = [
        "<!doctype html>",
        '<html lang="zh-Hans">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{html.escape(REPORT['title'])}</title>",
        f"<style>{css}</style>",
        "</head>",
        "<body><main class=\"page\">",
        "<section class=\"cover\">",
        f"<h1>{html.escape(REPORT['title'])}</h1>",
        f"<p class=\"subtitle\">{html.escape(REPORT['brand'])}｜{html.escape(REPORT['subtitle'])}</p>",
        "</section>",
        f"<p class=\"notice\">{html.escape(REPORT['notice'])}</p>",
        "<h2>诊断元信息</h2>",
        "<div class=\"meta-grid\">",
    ]
    for key, value in REPORT["meta"]:
        parts.append(f"<div>{html.escape(key)}</div><div>{html.escape(value)}</div>")
    parts.append("</div>")
    parts.append("<h2>执行摘要</h2><ul>")
    for title, detail in REPORT["summary"]:
        parts.append(f"<li><strong>{html.escape(title)}</strong>：{html.escape(detail)}</li>")
    parts.append("</ul>")

    section_specs = [
        ("方法依据与评分口径", "method_basis", REPORT["section_meanings"]["method_basis"]),
        ("问题样本覆盖矩阵", "question_coverage", REPORT["section_meanings"]["question_coverage"]),
        ("AI 答案可见性与竞品对比", "baseline", REPORT["section_meanings"]["baseline"]),
        ("竞品推荐差距", "competitors", REPORT["section_meanings"]["competitors"]),
        ("GEO 特征评分", "scores", REPORT["section_meanings"]["scores"]),
        ("机会地图与优先级矩阵", "opportunities", REPORT["section_meanings"]["opportunities"]),
        ("页面与内容资产修复清单", "repairs", ""),
        ("来源台账与待确认项", "sources", ""),
    ]
    for title, key, meaning in section_specs:
        if key == "sources":
            parts.append("<section class=\"avoid-break\">")
        parts.append(f"<h2>{html.escape(title)}</h2>")
        parts.append(render_html_table(HEADERS[key], REPORT[key]))
        if meaning:
            parts.append(f"<p class=\"meaning\">判断含义：{html.escape(meaning)}</p>")
        if key == "sources":
            parts.append("</section>")
    parts.append("<h2>问题样本</h2><ol>")
    for question in REPORT["sampling_questions"]:
        parts.append(f"<li>{html.escape(question)}</li>")
    parts.append("</ol>")
    parts.append("<h2>结论</h2>")
    parts.append(f"<p>{html.escape(REPORT['conclusion'])}</p>")
    parts.append("</main></body></html>")
    return "\n".join(parts)


def render_html_table(headers: list[str], rows: list[list[str]]) -> str:
    header_html = "".join(f"<th>{html.escape(item)}</th>" for item in headers)
    row_html = []
    for row in rows:
        row_html.append("<tr>" + "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row) + "</tr>")
    return f"<div class=\"table-wrap\"><table><thead><tr>{header_html}</tr></thead><tbody>{''.join(row_html)}</tbody></table></div>"


def set_cell_text(cell, text: str, bold: bool = False, shade: str | None = None, font_size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
    if shade:
        set_cell_shading(cell, shade)


def cm_to_twips(width_cm: float) -> int:
    return int(width_cm / 2.54 * 1440)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))


def set_fixed_table_layout(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = min(sum(widths_cm), DOCX_SAFE_TABLE_WIDTH_CM)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(cm_to_twips(table_width)))
    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), "0")
    tbl = table._tbl
    for existing in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(cm_to_twips(width)))
        grid.append(grid_col)
    tbl.insert(1, grid)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D8E0E8")


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 22, "18212F"),
        ("Heading 1", 15, "1D5D8C"),
        ("Heading 2", 12, "18212F"),
    ]:
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(7)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    font_size = 7.6 if len(headers) >= 8 else 8.0 if len(headers) >= 6 else 8.5
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, shade="F5F7FA", font_size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row):
            set_cell_text(cells[idx], cell_text, font_size=font_size)
    set_fixed_table_layout(table, column_widths(headers))
    doc.add_paragraph()


def column_widths(headers: list[str]) -> list[float]:
    presets = {
        tuple(HEADERS["method_basis"]): [2.2, 3.6, 5.0, 4.6],
        tuple(HEADERS["question_coverage"]): [1.6, 1.2, 1.7, 6.6, 4.0],
        tuple(HEADERS["baseline"]): [1.5, 1.7, 1.5, 1.5, 1.8, 1.8, 2.9, 2.6],
        tuple(HEADERS["competitors"]): [2.1, 1.6, 1.5, 1.5, 4.6, 4.2],
        tuple(HEADERS["scores"]): [2.0, 1.4, 7.3, 4.8],
        tuple(HEADERS["opportunities"]): [1.5, 1.8, 2.4, 1.1, 1.1, 1.1, 1.6, 3.6],
        tuple(HEADERS["repairs"]): [2.1, 3.4, 4.0, 2.8, 3.2],
        tuple(HEADERS["sources"]): [2.1, 2.8, 1.8, 2.9, 4.8],
        ("平台", "问题类型", "品牌是否出现", "是否被推荐", "引用源"): [2.0, 3.0, 2.2, 2.2, 5.6],
        ("平台", "高频竞品", "事实偏差", "备注"): [2.1, 3.8, 4.7, 4.6],
        ("机会类型", "机会", "对应问题", "价值", "工作量", "优先级"): [2.0, 3.0, 4.0, 1.4, 1.4, 1.4],
        ("机会", "负责人建议", "验收标准"): [3.5, 3.0, 8.7],
    }
    return presets.get(tuple(headers), [DOCX_SAFE_TABLE_WIDTH_CM / len(headers)] * len(headers))


def add_docx_named_table(doc: Document, key: str) -> None:
    headers = HEADERS[key]
    rows = REPORT[key]
    if key == "baseline":
        doc.add_paragraph("采样概览", style="Heading 2")
        add_table(
            doc,
            ["平台", "问题类型", "品牌是否出现", "是否被推荐", "引用源"],
            [[row[0], row[1], row[2], row[3], row[4]] for row in rows],
        )
        doc.add_paragraph("竞品、偏差与备注", style="Heading 2")
        add_table(
            doc,
            ["平台", "高频竞品", "事实偏差", "备注"],
            [[row[0], row[5], row[6], row[7]] for row in rows],
        )
        return
    if key == "opportunities":
        doc.add_paragraph("机会优先级", style="Heading 2")
        add_table(
            doc,
            ["机会类型", "机会", "对应问题", "价值", "工作量", "优先级"],
            [[row[0], row[1], row[2], row[3], row[4], row[5]] for row in rows],
        )
        doc.add_paragraph("责任人与验收", style="Heading 2")
        add_table(
            doc,
            ["机会", "负责人建议", "验收标准"],
            [[row[1], row[6], row[7]] for row in rows],
        )
        return
    add_table(doc, headers, rows)


def add_bullets(doc: Document, items: list[tuple[str, str]]) -> None:
    for title, detail in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{title}：")
        run.bold = True
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        paragraph.add_run(detail)


def render_docx() -> None:
    doc = Document()
    set_document_styles(doc)
    title = doc.add_paragraph(style="Title")
    title.add_run(REPORT["title"])
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"{REPORT['brand']}｜{REPORT['subtitle']}").bold = True
    notice = doc.add_paragraph()
    notice.add_run("说明：").bold = True
    notice.add_run(REPORT["notice"])

    doc.add_heading("诊断元信息", level=1)
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(meta_table)
    for key, value in REPORT["meta"]:
        cells = meta_table.add_row().cells
        set_cell_text(cells[0], key, bold=True, shade="F5F7FA")
        set_cell_text(cells[1], value)
    set_fixed_table_layout(meta_table, [3.0, 12.5])

    doc.add_heading("执行摘要", level=1)
    add_bullets(doc, REPORT["summary"])

    doc.add_heading("方法依据与评分口径", level=1)
    add_table(doc, HEADERS["method_basis"], REPORT["method_basis"])
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['method_basis']}")

    doc.add_heading("问题样本覆盖矩阵", level=1)
    add_table(doc, HEADERS["question_coverage"], REPORT["question_coverage"])
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['question_coverage']}")

    doc.add_heading("AI 答案可见性与竞品对比", level=1)
    add_docx_named_table(doc, "baseline")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['baseline']}")

    doc.add_heading("竞品推荐差距", level=1)
    add_docx_named_table(doc, "competitors")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['competitors']}")

    doc.add_heading("GEO 特征评分", level=1)
    add_docx_named_table(doc, "scores")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['scores']}")

    doc.add_heading("机会地图与优先级矩阵", level=1)
    add_docx_named_table(doc, "opportunities")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['opportunities']}")

    doc.add_heading("页面与内容资产修复清单", level=1)
    add_docx_named_table(doc, "repairs")

    doc.add_heading("问题样本", level=1)
    for question in REPORT["sampling_questions"]:
        doc.add_paragraph(question, style="List Number")

    doc.add_heading("来源台账与待确认项", level=1)
    add_docx_named_table(doc, "sources")

    doc.add_heading("结论", level=1)
    doc.add_paragraph(REPORT["conclusion"])
    doc.save(DOCX_PATH)


def render_pdf() -> None:
    chrome = Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
    if not chrome.exists():
        raise RuntimeError("未找到 Google Chrome，无法从 HTML 渲染 PDF。")
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            str(chrome),
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            "--disable-background-networking",
            "--disable-sync",
            "--disable-extensions",
            "--disable-component-update",
            "--disable-default-apps",
            f"--user-data-dir={tmpdir}",
            f"--print-to-pdf={PDF_PATH}",
            "--no-pdf-header-footer",
            "--print-to-pdf-no-header",
            HTML_PATH.as_uri(),
        ]
        process = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        try:
            process.wait(timeout=30)
        except subprocess.TimeoutExpired:
            process.terminate()
            try:
                process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                process.kill()
        if not PDF_PATH.exists() or PDF_PATH.stat().st_size <= 1024:
            raise RuntimeError("PDF 未成功生成或文件过小。")


def review() -> None:
    checks: list[tuple[str, bool, str]] = []
    for path in [MD_PATH, HTML_PATH, DOCX_PATH, PDF_PATH]:
        checks.append((path.name, path.exists() and path.stat().st_size > 1024, "文件存在且非空"))
    html_text = HTML_PATH.read_text(encoding="utf-8")
    md_text = MD_PATH.read_text(encoding="utf-8")
    docx_doc = Document(DOCX_PATH) if DOCX_PATH.exists() else None
    docx_tables = len(docx_doc.tables) if docx_doc else 0
    max_docx_table_width = 0.0
    if docx_doc:
        for table in docx_doc.tables:
            if not table.rows:
                continue
            width = sum((cell.width.cm if cell.width else 0.0) for cell in table.rows[0].cells)
            max_docx_table_width = max(max_docx_table_width, width)
    pdf_text = ""
    if shutil.which("pdftotext") and PDF_PATH.exists():
        pdf_text = subprocess.run(
            ["pdftotext", str(PDF_PATH), "-"],
            check=False,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            timeout=15,
        ).stdout
    checks.extend(
        [
            ("HTML 白底", "background: #ffffff" in html_text and "color-scheme: light" in html_text, "白底与浅色模式存在"),
            ("HTML 防溢出", "overflow-wrap: anywhere" in html_text and "table-layout: fixed" in html_text, "表格与长文本防溢出规则存在"),
            ("中文字段", all(token in md_text for token in ["诊断元信息", "执行摘要", "方法依据与评分口径", "机会地图", "验收标准"]), "核心字段为中文简体"),
            ("问题覆盖", all(token in md_text for token in ["推荐", "购买决策", "场景解决"]), "九类意图覆盖矩阵存在"),
            ("格式一致", all(token in html_text for token in ["方法依据与评分口径", "问题样本覆盖矩阵", "AI 答案可见性与竞品对比", "页面与内容资产修复清单", "问题样本"]), "HTML 与 Markdown 章节一致"),
            ("Word 表格结构", docx_tables >= 11, f"检测到 {docx_tables} 张表格"),
            ("Word 表格宽度", max_docx_table_width <= DOCX_SAFE_TABLE_WIDTH_CM + 0.1, f"最大表宽约 {max_docx_table_width:.1f}cm"),
            ("PDF 无本地路径页脚", "file:///" not in pdf_text and "Users/laoyao" not in pdf_text, "未检测到浏览器默认路径页脚"),
        ]
    )
    failed = [name for name, ok, _ in checks if not ok]
    lines = [
        "# 示例报告自检记录",
        "",
        "- 自检日期：2026-05-19",
        f"- 示例性质：{REPORT.get('example_nature', '合成样例，不代表真实平台长期结论')}",
        "- 版式基准：kami 专业文档纪律 + 白底报告约束",
        "",
        "| 检查项 | 结果 | 说明 |",
        "| --- | --- | --- |",
    ]
    for name, ok, detail in checks:
        lines.append(f"| {name} | {'通过' if ok else '需修复'} | {detail} |")
    lines.append("")
    if failed:
        lines.append("## 待修复")
        for item in failed:
            lines.append(f"- {item}")
        REVIEW_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
        raise RuntimeError("示例报告自检未通过：" + "、".join(failed))
    lines.append("## 自检结论")
    lines.append("四种格式均已生成；HTML/PDF 采用白底报告版式，表格设置固定布局与长文本换行；Word 使用标题样式、显式列宽、固定表格网格、表格边框和中文字段，密集 8 列表已按 Word 版拆成窄表。")
    lines.append("受本机缺少 LibreOffice 影响，未执行 DOCX 页面图片渲染检查；已完成结构与文件完整性检查。")
    REVIEW_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    MD_PATH.write_text(render_markdown(), encoding="utf-8")
    HTML_PATH.write_text(render_html(), encoding="utf-8")
    render_docx()
    render_pdf()
    review()


if __name__ == "__main__":
    main()
