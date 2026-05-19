#!/usr/bin/env python3
# Copyright © 2026 姚金刚. All rights reserved.
# Project: yao-geo-content-refiner
# Created by: 姚金刚
# Date: 2026-05-16
# X: https://x.com/yaojingang

from __future__ import annotations

import argparse, html, json
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

MODULES = ["报告摘要","原文 GEO 评分","GEO 改造版文章","改造前后差异报告","原子事实卡","FAQ 与同义问法","理论依据与改造映射","证据账本与缺口","页面发布版 HTML 建议","自 Review 结果"]
HEADERS = {
    "scorecard": ["维度","评分","证据判断","优先动作"],
    "diff_report": ["位置","原文问题","改造动作","改造后效果","风险"],
    "fact_cards": ["主体","属性","数值","时间","来源","适用边界","核验状态","可引用句"],
    "faq": ["问题","答案","类型"],
    "research_alignment": ["理论","研究启发","落地改造"],
    "sources": ["事实项","来源","证据状态","备注"],
    "self_review": ["检查项","检查结果","状态"],
}
DOCX_TABLE_WIDTH_DXA = 15260

def h(x: Any) -> str:
    return html.escape(str(x), quote=True)

def print_wrap(x: Any) -> str:
    text = str(x)
    if "http://" in text or "https://" in text:
        for token in ["/", ".", "?", "&", "=", "-"]:
            text = text.replace(token, token + " ")
    return text

def md_table(headers, rows) -> str:
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        out.append("| " + " | ".join(str(c).replace("|","｜").replace("\n","<br>") for c in row) + " |")
    return "\n".join(out)

def render_markdown(d):
    meta = "\n".join(f"- {k}：{v}" for k, v in d["meta"].items())
    summary = "\n".join(f"- {x}" for x in d["summary"])
    article = []
    for sec in d["refined_article"]:
        article.append(f"### {sec['heading']}")
        article += sec.get("paragraphs", [])
        if "table" in sec:
            article.append(md_table(sec["table"]["headers"], sec["table"]["rows"]))
    cms = "\n".join(f"{i}. {x}" for i, x in enumerate(d["cms_html_advice"], 1))
    return f"""# {d['title']}

**{d['subtitle']}**

> {d['notice']}

## 元信息

{meta}

## 报告摘要

{summary}

## 原文 GEO 评分

{md_table(HEADERS['scorecard'], d['scorecard'])}

## GEO 改造版文章

{chr(10).join(article)}

## 改造前后差异报告

{md_table(HEADERS['diff_report'], d['diff_report'])}

## 原子事实卡

{md_table(HEADERS['fact_cards'], d['fact_cards'])}

## FAQ 与同义问法

{md_table(HEADERS['faq'], d['faq'])}

## 理论依据与改造映射

{md_table(HEADERS['research_alignment'], d['research_alignment'])}

## 证据账本与缺口

{md_table(HEADERS['sources'], d['sources'])}

## 页面发布版 HTML 建议

{cms}

## 自 Review 结果

{md_table(HEADERS['self_review'], d['self_review'])}
"""

def html_table(headers, rows):
    th = "".join(f"<th>{h(x)}</th>" for x in headers)
    trs = "".join("<tr>" + "".join(f"<td>{h(c)}</td>" for c in r) + "</tr>" for r in rows)
    return f"<div class='table-wrap'><table><thead><tr>{th}</tr></thead><tbody>{trs}</tbody></table></div>"

def render_html(d):
    article = []
    for sec in d["refined_article"]:
        article.append(f"<h3>{h(sec['heading'])}</h3>")
        article += [f"<p>{h(p)}</p>" for p in sec.get("paragraphs", [])]
        if "table" in sec:
            article.append(html_table(sec["table"]["headers"], sec["table"]["rows"]))
    return f"""<!doctype html><html lang="zh-CN"><head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><title>{h(d['title'])}</title>
<style>
body{{margin:0;background:#ffffff;color:#172033;font-family:-apple-system,BlinkMacSystemFont,"PingFang SC","Microsoft YaHei",sans-serif;line-height:1.68;letter-spacing:0}}
main{{width:min(980px,calc(100vw - 40px));margin:0 auto;padding:42px 0 56px}} header{{border-bottom:2px solid #172033;padding-bottom:22px;margin-bottom:28px}}
h1{{font-size:30px;line-height:1.25;margin:0 0 10px}} h2{{font-size:21px;border-bottom:1px solid #d9e0e8;padding-bottom:6px;margin:34px 0 14px}} h3{{font-size:17px;margin:22px 0 10px}}
.notice{{padding:12px 14px;border:1px solid #d9e0e8;background:#f6f8fb}} .table-wrap{{width:100%;overflow-x:auto;margin:12px 0 18px}}
table{{width:100%;border-collapse:collapse;table-layout:fixed;background:#fff;font-size:13px}} th,td{{border:1px solid #d9e0e8;padding:8px 9px;text-align:left;vertical-align:top;overflow-wrap:anywhere;word-break:break-word}} th{{background:#eef3f8;color:#1c334f}}
</style></head><body><main><header><h1>{h(d['title'])}</h1><p>{h(d['subtitle'])}</p><p class="notice">{h(d['notice'])}</p></header>
<h2>元信息</h2>{html_table(["字段","内容"], [[k,v] for k,v in d["meta"].items()])}
<h2>报告摘要</h2><ul>{''.join(f'<li>{h(x)}</li>' for x in d['summary'])}</ul>
<h2>原文 GEO 评分</h2>{html_table(HEADERS['scorecard'], d['scorecard'])}
<h2>GEO 改造版文章</h2>{''.join(article)}
<h2>改造前后差异报告</h2>{html_table(HEADERS['diff_report'], d['diff_report'])}
<h2>原子事实卡</h2>{html_table(HEADERS['fact_cards'], d['fact_cards'])}
<h2>FAQ 与同义问法</h2>{html_table(HEADERS['faq'], d['faq'])}
<h2>理论依据与改造映射</h2>{html_table(HEADERS['research_alignment'], d['research_alignment'])}
<h2>证据账本与缺口</h2>{html_table(HEADERS['sources'], d['sources'])}
<h2>页面发布版 HTML 建议</h2><ol>{''.join(f'<li>{h(x)}</li>' for x in d['cms_html_advice'])}</ol>
<h2>自 Review 结果</h2>{html_table(HEADERS['self_review'], d['self_review'])}</main></body></html>"""

def set_rfont(run, size=9.0, bold=None):
    run.font.name = "宋体"; run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts"); rpr.append(rf)
    rf.set(qn("w:eastAsia"), "宋体")

def cell_margins(cell):
    pr = cell._tc.get_or_add_tcPr()
    mar = pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
    if mar.getparent() is None: pr.append(mar)
    for name, val in {"top":120,"start":110,"bottom":120,"end":110}.items():
        node = mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
        if node.getparent() is None: mar.append(node)
        node.set(qn("w:w"), str(val)); node.set(qn("w:type"), "dxa")

def shade(cell, color="EEF3F8"):
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), color); cell._tc.get_or_add_tcPr().append(shd)

def widths(headers):
    weights = []
    for x in headers:
        if x in {"评分","风险","类型","状态","时间","格式","核验状态"}: weights.append(.75)
        elif x in {"问题","答案","证据判断","优先动作","备注","可引用句"}: weights.append(1.8)
        else: weights.append(1.2)
    total = sum(weights)
    out = [max(900, int(DOCX_TABLE_WIDTH_DXA * w / total)) for w in weights]
    out[-1] += DOCX_TABLE_WIDTH_DXA - sum(out)
    return out

def set_table(table, ws, font=8.4):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.autofit = False
    tbl = table._tbl; pr = tbl.tblPr
    tw = pr.find(qn("w:tblW"))
    if tw is None:
        tw = OxmlElement("w:tblW")
    if tw.getparent() is None: pr.append(tw)
    tw.set(qn("w:type"), "dxa"); tw.set(qn("w:w"), str(sum(ws)))
    layout = pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
    if layout.getparent() is None: pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
    if grid.getparent() is None: tbl.insert(0, grid)
    for child in list(grid): grid.remove(child)
    for w in ws:
        gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); grid.append(gc)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP; cell_margins(cell)
            cpr = cell._tc.get_or_add_tcPr()
            tcw = cpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
            if tcw.getparent() is None: cpr.append(tcw)
            tcw.set(qn("w:type"), "dxa"); tcw.set(qn("w:w"), str(ws[c]))
            if r == 0: shade(cell)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.12; p.paragraph_format.space_after = Pt(0)
                for run in p.runs: set_rfont(run, font, r == 0)

def add_table(doc, headers, rows):
    if len(headers) > 4:
        for i, row in enumerate(rows, 1):
            p = doc.add_paragraph(f"记录 {i}"); p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(3)
            for run in p.runs: set_rfont(run, 9.5, True)
            t = doc.add_table(rows=len(headers), cols=2)
            for r, head in enumerate(headers):
                t.rows[r].cells[0].text = head
                t.rows[r].cells[1].text = print_wrap(row[r] if r < len(row) else "")
            set_table(t, [2500, DOCX_TABLE_WIDTH_DXA - 2500], 8.8)
        return
    t = doc.add_table(rows=1, cols=len(headers))
    for i, x in enumerate(headers): t.rows[0].cells[i].text = x
    for row in rows:
        cs = t.add_row().cells
        for i, x in enumerate(row): cs[i].text = print_wrap(x)
    set_table(t, widths(headers), 8.4)

def heading(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    for run in p.runs: set_rfont(run, 15 if level == 2 else 11.5, True)

def render_docx(d, path):
    doc = Document(); sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE; sec.page_width = Cm(29.7); sec.page_height = Cm(21.0)
    sec.top_margin = sec.bottom_margin = Cm(1.25); sec.left_margin = sec.right_margin = Cm(1.35)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.font.name = "宋体"
    rpr = normal._element.get_or_add_rPr()
    rf = rpr.rFonts
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), "宋体")
    title = doc.add_heading(d["title"], 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs: set_rfont(run, 18, True)
    for text in [d["subtitle"], d["notice"]]:
        p = doc.add_paragraph(text)
        for run in p.runs: set_rfont(run, 10)
    heading(doc, "元信息"); add_table(doc, ["字段","内容"], [[k,v] for k,v in d["meta"].items()])
    heading(doc, "报告摘要")
    for x in d["summary"]:
        p = doc.add_paragraph(x, style="List Bullet")
        for run in p.runs: set_rfont(run, 10)
    heading(doc, "原文 GEO 评分"); add_table(doc, HEADERS["scorecard"], d["scorecard"])
    heading(doc, "GEO 改造版文章")
    for sec_data in d["refined_article"]:
        heading(doc, sec_data["heading"], 3)
        for para in sec_data.get("paragraphs", []):
            p = doc.add_paragraph(para)
            for run in p.runs: set_rfont(run, 10)
        if "table" in sec_data: add_table(doc, sec_data["table"]["headers"], sec_data["table"]["rows"])
    for key, title in [("diff_report","改造前后差异报告"),("fact_cards","原子事实卡"),("faq","FAQ 与同义问法"),("research_alignment","理论依据与改造映射"),("sources","证据账本与缺口")]:
        heading(doc, title); add_table(doc, HEADERS[key], d[key])
    heading(doc, "页面发布版 HTML 建议")
    for x in d["cms_html_advice"]:
        p = doc.add_paragraph(x, style="List Number")
        for run in p.runs: set_rfont(run, 10)
    heading(doc, "自 Review 结果"); add_table(doc, HEADERS["self_review"], d["self_review"])
    doc.save(path)

def pstyle():
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light")); s = getSampleStyleSheet()
    return {
        "title": ParagraphStyle("t", parent=s["Title"], fontName="STSong-Light", fontSize=19, leading=25, alignment=TA_CENTER),
        "h2": ParagraphStyle("h2", parent=s["Heading2"], fontName="STSong-Light", fontSize=14, leading=18, spaceBefore=11, spaceAfter=6),
        "h3": ParagraphStyle("h3", parent=s["Heading3"], fontName="STSong-Light", fontSize=11.5, leading=16, spaceBefore=8, spaceAfter=4),
        "p": ParagraphStyle("p", parent=s["Normal"], fontName="STSong-Light", fontSize=9.4, leading=14.3, spaceAfter=5),
        "td": ParagraphStyle("td", parent=s["Normal"], fontName="STSong-Light", fontSize=7.3, leading=9.8),
        "th": ParagraphStyle("th", parent=s["Normal"], fontName="STSong-Light", fontSize=7.6, leading=10, textColor=colors.HexColor("#1c334f")),
    }

def pp(x, style): return Paragraph(h(print_wrap(x)).replace("\n","<br/>"), style)

def add_pdf_table(story, headers, rows, st):
    page_w = A4[0] - 3.0 * cm
    if len(headers) > 4:
        for i, row in enumerate(rows, 1):
            story.append(pp(f"记录 {i}", st["h3"]))
            data = [[pp(head, st["th"]), pp(row[j] if j < len(row) else "", st["td"])] for j, head in enumerate(headers)]
            t = Table(data, colWidths=[3.0*cm, page_w-3.0*cm], hAlign="LEFT")
            t.setStyle(TableStyle([("BACKGROUND",(0,0),(0,-1),colors.HexColor("#eef3f8")),("GRID",(0,0),(-1,-1),.45,colors.HexColor("#d9e0e8")),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),4),("BOTTOMPADDING",(0,0),(-1,-1),4)]))
            story += [t, Spacer(1, .16*cm)]
        return
    weights = [0.75 if x in {"评分","风险","类型","状态","时间"} else 1.5 for x in headers]
    tw = sum(weights); cols = [page_w*w/tw for w in weights]
    data = [[pp(x, st["th"]) for x in headers]] + [[pp(x, st["td"]) for x in row] for row in rows]
    t = Table(data, colWidths=cols, repeatRows=1, hAlign="LEFT")
    t.setStyle(TableStyle([("BACKGROUND",(0,0),(-1,0),colors.HexColor("#eef3f8")),("GRID",(0,0),(-1,-1),.45,colors.HexColor("#d9e0e8")),("VALIGN",(0,0),(-1,-1),"TOP"),("LEFTPADDING",(0,0),(-1,-1),4),("RIGHTPADDING",(0,0),(-1,-1),4),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    story += [t, Spacer(1, .22*cm)]

def render_pdf(d, path):
    st = pstyle(); story = [pp(d["title"], st["title"]), pp(d["subtitle"], st["p"]), pp(d["notice"], st["p"])]
    story.append(pp("元信息", st["h2"])); add_pdf_table(story, ["字段","内容"], [[k,v] for k,v in d["meta"].items()], st)
    story.append(pp("报告摘要", st["h2"]))
    for x in d["summary"]: story.append(pp("· " + x, st["p"]))
    story.append(pp("原文 GEO 评分", st["h2"])); add_pdf_table(story, HEADERS["scorecard"], d["scorecard"], st)
    story.append(pp("GEO 改造版文章", st["h2"]))
    for sec in d["refined_article"]:
        story.append(pp(sec["heading"], st["h3"]))
        for para in sec.get("paragraphs", []): story.append(pp(para, st["p"]))
        if "table" in sec: add_pdf_table(story, sec["table"]["headers"], sec["table"]["rows"], st)
    for key, title in [("diff_report","改造前后差异报告"),("fact_cards","原子事实卡"),("faq","FAQ 与同义问法"),("research_alignment","理论依据与改造映射"),("sources","证据账本与缺口")]:
        story.append(pp(title, st["h2"])); add_pdf_table(story, HEADERS[key], d[key], st)
    story.append(pp("页面发布版 HTML 建议", st["h2"]))
    for i, x in enumerate(d["cms_html_advice"], 1): story.append(pp(f"{i}. {x}", st["p"]))
    story.append(pp("自 Review 结果", st["h2"])); add_pdf_table(story, HEADERS["self_review"], d["self_review"], st)
    SimpleDocTemplate(str(path), pagesize=A4, leftMargin=1.5*cm, rightMargin=1.5*cm, topMargin=1.5*cm, bottomMargin=1.5*cm).build(story)

def inspect_docx_layout(path):
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    sects = root.findall(".//w:sectPr", ns)
    if not sects:
        return {"status": "needs-fix", "problems": ["DOCX has no section properties."]}
    sect = sects[-1]
    pg = sect.find("w:pgSz", ns)
    mar = sect.find("w:pgMar", ns)
    if pg is None or mar is None:
        return {"status": "needs-fix", "problems": ["DOCX page size or margins missing."]}
    w_ns = "{" + ns["w"] + "}"
    page_width = int(pg.attrib[w_ns + "w"])
    usable_width = page_width - int(mar.attrib[w_ns + "left"]) - int(mar.attrib[w_ns + "right"])
    tables = []
    bad = []
    for idx, tbl in enumerate(root.findall(".//w:tbl", ns), 1):
        cols = []
        for col in tbl.findall("./w:tblGrid/w:gridCol", ns):
            width = col.attrib.get(w_ns + "w")
            if width:
                cols.append(int(width))
        total = sum(cols)
        item = {"index": idx, "columns": len(cols), "width_dxa": total, "right_margin_dxa": usable_width - total}
        tables.append(item)
        if total > usable_width:
            bad.append(item)
    return {
        "status": "pass" if not bad else "needs-fix",
        "orientation": pg.attrib.get(w_ns + "orient", "portrait"),
        "usable_width_dxa": usable_width,
        "table_count": len(tables),
        "max_table_width_dxa": max((x["width_dxa"] for x in tables), default=0),
        "min_right_margin_dxa": min((x["right_margin_dxa"] for x in tables), default=usable_width),
        "overflow_tables": bad,
    }

def write_quality(d, paths, outdir):
    md = paths["markdown"].read_text(encoding="utf-8"); ht = paths["html"].read_text(encoding="utf-8")
    problems = []
    checks = []
    for k, p in paths.items():
        checks.append({"format": k, "path": str(p), "exists": p.exists(), "bytes": p.stat().st_size if p.exists() else 0})
        if not p.exists() or p.stat().st_size < 512: problems.append(f"{k} missing or too small")
    module_checks = [{"module": m, "in_markdown": m in md, "in_html": m in ht} for m in MODULES]
    problems += [f"module missing: {x['module']}" for x in module_checks if not x["in_markdown"] or not x["in_html"]]
    if "/Users/" in ht:
        problems.append("html contains local absolute path")
    docx_layout = inspect_docx_layout(paths["docx"])
    if docx_layout["status"] != "pass":
        problems.append("docx table width exceeds page usable width")
    (outdir/"quality-report.json").write_text(json.dumps({"title": d["title"], "checks": checks, "module_checks": module_checks, "layout_profile": "fixed-width-docx-wide-table-cards", "layout_logic": {"docx_orientation": "landscape", "docx_table_width_dxa": DOCX_TABLE_WIDTH_DXA, "wide_table_rule": "Word/PDF 中超过 4 列的表格自动转纵向事实卡", "url_wrapping": "Word/PDF 中长 URL 强制可换行"}, "docx_layout": docx_layout, "problems": problems, "status": "pass" if not problems else "needs-fix"}, ensure_ascii=False, indent=2), encoding="utf-8")

def main():
    ap = argparse.ArgumentParser(); ap.add_argument("--input", type=Path, required=True); ap.add_argument("--output-dir", type=Path, required=True); ap.add_argument("--basename", required=True)
    a = ap.parse_args(); d = json.loads(a.input.read_text(encoding="utf-8")); a.output_dir.mkdir(parents=True, exist_ok=True)
    paths = {"markdown": a.output_dir/f"{a.basename}.md", "html": a.output_dir/f"{a.basename}.html", "docx": a.output_dir/f"{a.basename}.docx", "pdf": a.output_dir/f"{a.basename}.pdf"}
    paths["markdown"].write_text(render_markdown(d), encoding="utf-8"); paths["html"].write_text(render_html(d), encoding="utf-8")
    render_docx(d, paths["docx"]); render_pdf(d, paths["pdf"]); write_quality(d, paths, a.output_dir)

if __name__ == "__main__":
    main()
